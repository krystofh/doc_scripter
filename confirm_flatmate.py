import sys
import json
import os
from docx import Document


def display_help():
    print(
        "Usage: python confirm_flatmate.py <word_document_filename> <config_json_file>"
    )
    sys.exit(1)


def load_config(config_file):
    try:
        with open(config_file, "r") as file:
            config_data = json.load(file)
        return config_data
    except FileNotFoundError:
        print(f"Error: Config file '{config_file}' not found.")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON format in config file '{config_file}'.")
        sys.exit(1)


def print_rows(doc, table_number: int):
    """Print all rows of a table specified by its index. Useful for playing around in order to find what element you are looking for.

    Parameters
    ----------
    doc : docx.Document
        Word document to be parsed
    table_number : int
        index of the table
    """
    for i, row in enumerate(doc.tables[table_number].rows):
        print(f"Row {i}: ", end="")
        print(row.cells[0].text)


def check_keyword(config_data: dict, key: str) -> int:
    for obj, properties in config_data.items():
        for obj_property in properties:
            if config_data[obj][obj_property]["keyword"] in key: # handle cases when token like 'SURNAME,' or 'ADDRESS:'
                value = config_data[obj][obj_property]["value"]
                return value
    return ''

def process_paragraph(text: str, config_data: dict):
    tokens = text.split(' ')
    for i, token in enumerate(tokens):
        new_token = check_keyword(config_data, token)  # check if token is a keyword and if so, replace it
        if new_token != '':
            print(f"Replacing {token} with: {new_token}")
            tokens[i] = new_token
    return ' '.join(tokens) # return entire, joined, paragraph

def replace_keywords(doc, config_data: dict, mode="table"):
    if mode == "paragraph":
        pass
    elif mode == "table":
        # Used if the document has a table structure
        # Look up each data keyword to be changed in the document (see config.json)
        # "obj": {"property1": {"keyword": KEY, "value": VALUE}, "property2": ...}
        # Example object: "main_tenant": {"firstname": {"keyword": "FIRSTNAME_MAIN_TENANT","value": "Max"}

        # Look in each table
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Each cell can have multiple paragraphs with different formatting
                    # In order not to overwrite formatting, edit the paragraphs, not entire cell
                    for cell_paragraph in cell.paragraphs:
                        new_text = process_paragraph(cell_paragraph.text, config_data)
                        cell_paragraph.text = new_text # rewrite the text
                        
    else:
        raise AttributeError("Unsopported operation mode for keyword search")


def main():
    # Check for the correct number of arguments
    if len(sys.argv) != 3:
        display_help()

    # Get file names from command line arguments
    docx_filename = sys.argv[1]
    json_filename = sys.argv[2]

    # Check if files exist
    if not os.path.exists(docx_filename):
        print(f"Error: Word document file '{docx_filename}' not found.")
        sys.exit(1)

    config_data = load_config(json_filename)

    # Load the Word document
    document = Document(docx_filename)
    # print_rows(document, 0)

    # Replace keywords in the document
    replace_keywords(document, config_data)

    # Save the modified document
    output_filename = f"modified_{docx_filename}"
    document.save(output_filename)
    print(f"Success: Document saved as '{output_filename}'.")


if __name__ == "__main__":
    main()
